def parse_document(url: str) -> list:
    """
    Wrapper for backward compatibility. Calls parse_document_url.
    """
    return parse_document_url(url)
# app/parser.py

import os
import requests
import fitz  # PyMuPDF
import pdfplumber
import docx
import tempfile
from io import BytesIO
from pathlib import Path
from typing import List, Dict
from app.utils import chunk_text, detect_file_type

async def process_document(url: str) -> str:
    """
    Downloads the document from a given URL and saves locally.
    Returns the local file path.
    """
    response = requests.get(url)
    response.raise_for_status()

    filename = str(url).split("/")[-1].split("?")[0]
    ext = filename.split(".")[-1].lower()

    temp_dir = Path("data/docs")
    temp_dir.mkdir(parents=True, exist_ok=True)
    file_path = temp_dir / f"{filename}"

    with open(file_path, "wb") as f:
        f.write(response.content)

    return str(file_path)


async def extract_text_chunks(file_path: str) -> List[Dict]:
    """
    Extracts text and returns a list of chunk dicts:
    [{ 'text': ..., 'metadata': {...} }, ...]
    """
    file_type = detect_file_type(file_path)

    if file_type == "pdf":
        return parse_pdf(file_path)
    elif file_type == "docx":
        return parse_docx(file_path)
    else:
        raise ValueError("Unsupported file type")


def parse_pdf(file_path: str) -> List[Dict]:
    """
    Parses PDF and chunks by page or paragraph
    """
    doc = fitz.open(file_path)
    chunks = []

    for page_number in range(len(doc)):
        page = doc.load_page(page_number)
        text = page.get_text().strip()
        split_chunks = chunk_text(text)

        for chunk in split_chunks:
            chunks.append({
                "text": chunk,
                "metadata": {
                    "page": page_number + 1,
                    "source": Path(file_path).name
                }
            })
    return chunks


def parse_docx(file_path: str) -> List[Dict]:
    """
    Parses DOCX and chunks by paragraph
    """
    doc = docx.Document(file_path)
    chunks = []

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            split_chunks = chunk_text(text)
            for chunk in split_chunks:
                chunks.append({
                    "text": chunk,
                    "metadata": {
                        "paragraph": i + 1,
                        "source": Path(file_path).name
                    }
                })
    return chunks


def parse_pdf_url(url: str) -> list:
    response = requests.get(url)
    with pdfplumber.open(BytesIO(response.content)) as pdf:
        text = ''
        for page in pdf.pages:
            text += page.extract_text() + '\n'
    return chunk_text(text)


def parse_docx_url(url: str) -> list:
    response = requests.get(url)
    doc = docx.Document(BytesIO(response.content))
    text = '\n'.join([para.text for para in doc.paragraphs])
    return chunk_text(text)


def parse_document_url(url: str) -> list:
    if url.lower().endswith('.pdf'):
        return parse_pdf_url(url)
    elif url.lower().endswith('.docx'):
        return parse_docx_url(url)
    else:
        raise ValueError('Unsupported document type')
